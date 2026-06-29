"""
Barcha xodimlarning login/parol ma'lumotlarini login_parol.docx (haqiqiy .docx) faylga yozadi.
Parol = telefon raqami (+ belgisiz).
Ishlatish: python export_logins.py
"""
import sys, os

# DB_URL: agar env'da bo'lsa o'shani, bo'lmasa default (Docker DB 5432-portda)
os.environ.setdefault("DATABASE_URL", "postgresql://crm_user:crm_pass@127.0.0.1:5432/crm_db")
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.database import SessionLocal
from app import models

ROLE_LABEL = {
    "superadmin":         "Superadmin",
    "direktor":           "Direktor",
    "zamdirektor":        "Zamdirektor",
    "bolim_boshligi":     "Bo'lim boshlig'i",
    "boshqarma_boshligi": "Boshqarma boshlig'i",
    "xodim":              "Xodim",
}


def set_cell_bg(cell, color_hex):
    """Hujayra fon rangini o'rnatadi."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def set_cell_text(cell, text, *, bold=False, size=9, color=None, align="left"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def main():
    db = SessionLocal()
    try:
        depts = {d.id: d.name for d in db.query(models.Department).all()}
        emps = (
            db.query(models.Employee)
            .order_by(models.Employee.department_id.is_(None).desc(),
                      models.Employee.department_id,
                      models.Employee.id)
            .all()
        )

        doc = Document()

        # Sahifa — kichik margin (ko'p ustun sig'ishi uchun)
        for section in doc.sections:
            section.left_margin   = Inches(0.4)
            section.right_margin  = Inches(0.4)
            section.top_margin    = Inches(0.5)
            section.bottom_margin = Inches(0.5)

        # ── Sarlavha ──
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = title.add_run("TMSITI CRM — Tizimga kirish ma'lumotlari")
        r.bold = True
        r.font.size = Pt(15)
        r.font.color.rgb = RGBColor(0x04, 0x16, 0x38)

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rs = sub.add_run(
            f"Texnik me'yorlash va standartlashtirish ilmiy-tadqiqot instituti  ·  "
            f"Jami {len(emps)} ta foydalanuvchi"
        )
        rs.font.size = Pt(9)
        rs.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()  # bo'sh joy

        # ── Jadval ──
        headers = ["№", "F.I.Sh.", "Bo'lim / Boshqarma", "Lavozim", "Telefon (login)", "Parol", "Rol"]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Ustun kengliklari
        widths = [Inches(0.35), Inches(1.6), Inches(2.2), Inches(1.5), Inches(1.3), Inches(1.1), Inches(1.1)]

        # Header row
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            set_cell_text(hdr[i], h, bold=True, size=9,
                          color=(0xFF, 0xFF, 0xFF),
                          align="center" if i not in (1, 2, 3) else "left")
            set_cell_bg(hdr[i], "041638")

        # Data rows
        for idx, e in enumerate(emps, 1):
            dept_name = depts.get(e.department_id, "—") if e.department_id else "—"
            password  = e.phone.lstrip("+")
            role_key  = e.role.value if hasattr(e.role, "value") else str(e.role)
            role      = ROLE_LABEL.get(role_key, role_key)
            bg        = "FFFFFF" if idx % 2 else "F2F5FA"

            cells = table.add_row().cells
            set_cell_text(cells[0], idx, align="center", size=9)
            set_cell_text(cells[1], e.full_name, bold=True, size=9)
            set_cell_text(cells[2], dept_name, size=8)
            set_cell_text(cells[3], e.position, size=8)
            set_cell_text(cells[4], e.phone, bold=True, align="center", size=9, color=(0x1D, 0x4E, 0xD8))
            set_cell_text(cells[5], password, bold=True, align="center", size=9, color=(0x99, 0x1B, 0x1B))
            set_cell_text(cells[6], role, align="center", size=8)

            for c in cells:
                set_cell_bg(c, bg)
            # widthlar
            for c, w in zip(cells, widths):
                c.width = w

        # Header width
        for c, w in zip(hdr, widths):
            c.width = w

        # ── Izoh ──
        doc.add_paragraph()
        note = doc.add_paragraph()
        rn1 = note.add_run("Eslatma: ")
        rn1.bold = True
        rn1.font.size = Pt(9)
        rn2 = note.add_run(
            "Parol — telefon raqamining \"+\" belgisiz ko'rinishi "
            "(masalan: telefon +998901000070 → parol 998901000070). "
            "Birinchi kirishdan so'ng parolni o'zgartirish tavsiya etiladi."
        )
        rn2.font.size = Pt(9)
        rn2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        out_path = os.path.join(os.path.dirname(__file__), "login_parol.docx")
        doc.save(out_path)
        print(f"[OK] {out_path}")
        print(f"[OK] Jami {len(emps)} ta foydalanuvchi yozildi")

    finally:
        db.close()


if __name__ == "__main__":
    main()
